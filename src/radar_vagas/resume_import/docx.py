from __future__ import annotations

import re
from collections.abc import Iterator
from io import BytesIO
from typing import Any

from radar_vagas.domain.enums import ExtractedBlockType
from radar_vagas.domain.errors import RadarError
from radar_vagas.resume_import.models import ExtractedBlock, ExtractedDocument
from radar_vagas.resume_import.security import MAX_EXTRACTED_TEXT_CHARS, MIN_TEXT_CHARS


def extract_docx_document(content: bytes) -> ExtractedDocument:
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise RadarError("Instale o extra web para importar curriculos DOCX.") from exc

    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise RadarError("Nao foi possivel ler este DOCX.") from exc

    blocks: list[ExtractedBlock] = []
    order = 0
    total_chars = 0
    table_index = 0
    for element in _iter_body_items(document, Paragraph=Paragraph, Table=Table):
        if isinstance(element, Paragraph):
            line = _clean_text(element.text)
            if not line:
                continue
            block_type = _paragraph_block_type(element, line)
            heading = line.rstrip(":") if block_type == ExtractedBlockType.HEADING else None
            blocks.append(
                ExtractedBlock(
                    block_id=f"d-b{order + 1}",
                    order=order,
                    text=line,
                    page_number=None,
                    block_type=block_type,
                    heading=heading,
                )
            )
            order += 1
            total_chars += len(line)
        else:
            table_index += 1
            for row_index, row in enumerate(element.rows):
                seen_in_row: set[str] = set()
                for cell_index, cell in enumerate(row.cells):
                    cell_text = _clean_text("\n".join(p.text for p in cell.paragraphs))
                    if not cell_text or cell_text in seen_in_row:
                        continue
                    seen_in_row.add(cell_text)
                    blocks.append(
                        ExtractedBlock(
                            block_id=f"d-t{table_index}-r{row_index + 1}-c{cell_index + 1}",
                            order=order,
                            text=cell_text,
                            page_number=None,
                            block_type=ExtractedBlockType.TABLE_CELL,
                            table_index=table_index,
                            row_index=row_index + 1,
                            cell_index=cell_index + 1,
                        )
                    )
                    order += 1
                    total_chars += len(cell_text)
        if total_chars > MAX_EXTRACTED_TEXT_CHARS:
            raise RadarError("Texto extraido grande demais para importacao local segura.")

    if total_chars < MIN_TEXT_CHARS:
        raise RadarError("O DOCX tem pouca informacao textual para montar um perfil revisavel.")
    return ExtractedDocument(
        blocks=tuple(blocks),
        warnings=(),
        page_count=1,
        source_format="docx",
        extracted_character_count=total_chars,
        quality="textual",
    )


def _iter_body_items(document: Any, *, Paragraph: type[Any], Table: type[Any]) -> Iterator[Any]:
    parent_elm = document.element.body
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def _paragraph_block_type(paragraph: Any, line: str) -> ExtractedBlockType:
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    normalized_style = style_name.lower()
    if normalized_style.startswith(("heading", "titulo", "title")):
        return ExtractedBlockType.HEADING
    if line.endswith(":") and len(line) <= 80:
        return ExtractedBlockType.HEADING
    if re.match(r"^[-*•]\s+", line) or re.match(r"^\d+[.)]\s+", line):
        return ExtractedBlockType.LIST_ITEM
    return ExtractedBlockType.PARAGRAPH


def _clean_text(text: str) -> str:
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()
