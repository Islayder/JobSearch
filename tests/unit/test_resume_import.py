from __future__ import annotations

import zipfile
from collections.abc import Sequence
from io import BytesIO
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject
from sqlalchemy import select

from radar_vagas.config.settings import Settings
from radar_vagas.domain.enums import (
    ResumeImportCandidateType,
    ResumeImportConfidenceLabel,
    ResumeImportDecision,
    ResumeImportStatus,
)
from radar_vagas.domain.errors import RadarError
from radar_vagas.persistence.database import session_scope
from radar_vagas.persistence.migrations import run_migrations
from radar_vagas.persistence.models import ProfessionalProfileVersion, ResumeImportSession
from radar_vagas.resume_import.extraction import extract_resume
from radar_vagas.resume_import.parser import parse_resume_document
from radar_vagas.resume_import.pdf import extract_pdf_document
from radar_vagas.resume_import.security import validate_resume_upload
from radar_vagas.resume_import.service import (
    accept_candidate,
    confirm_import,
    create_import_session,
    discard_import,
    get_import_session,
    retry_import_session,
)


def test_security_rejects_old_doc_and_external_docx_relationship() -> None:
    with pytest.raises(RadarError, match="\\.doc antigo"):
        validate_resume_upload("curriculo.doc", b"legacy")

    content = BytesIO()
    with zipfile.ZipFile(content, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("word/document.xml", "<w:document/>")
        archive.writestr(
            "word/_rels/document.xml.rels",
            '<Relationships><Relationship TargetMode="External" Target="https://example.test"/></Relationships>',
        )

    with pytest.raises(RadarError, match="referencias externas"):
        validate_resume_upload("curriculo.docx", content.getvalue())


def test_pdf_without_text_gets_human_scanned_message() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    content = BytesIO()
    writer.write(content)

    with pytest.raises(
        RadarError, match="Não foi possível encontrar texto suficiente neste PDF\\."
    ):
        extract_resume("curriculo.pdf", content.getvalue())


def test_resume_upload_preserves_safe_unicode_filename() -> None:
    upload = validate_resume_upload("C:\\temp\\Currículo.pdf", b"%PDF-1.4\n")

    assert upload.filename == "Currículo.pdf"


def test_pdf_automatic_prefers_plain_when_layout_is_glued(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _patch_pdf_reader(
        monkeypatch,
        [
            _FakePdfPage(
                plain=_good_resume_text(),
                layout=_glued_resume_text(),
                fragments=(),
            )
        ],
    )

    document = extract_pdf_document(_fake_pdf_bytes())

    assert document.quality in {"GOOD", "ACCEPTABLE"}
    assert document.quality_metrics["selected_modes"] == ["plain"]
    assert any("Texto normal" in warning for warning in document.warnings)


def test_pdf_automatic_prefers_layout_when_plain_is_glued(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _patch_pdf_reader(
        monkeypatch,
        [
            _FakePdfPage(
                plain=_glued_resume_text(),
                layout=_good_resume_text(),
                fragments=(),
            )
        ],
    )

    document = extract_pdf_document(_fake_pdf_bytes())

    assert document.quality in {"GOOD", "ACCEPTABLE"}
    assert document.quality_metrics["selected_modes"] == ["layout"]
    assert any("Layout" in warning for warning in document.warnings)


def test_pdf_geometric_fallback_rebuilds_synthetic_resume_and_parser_items(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _patch_pdf_reader(
        monkeypatch,
        [
            _FakePdfPage(
                plain=_glued_resume_text(),
                layout=_glued_resume_text(),
                fragments=_resume_fragments(),
            )
        ],
    )

    _upload, document = extract_resume("curriculo.pdf", _fake_pdf_bytes())
    candidates = parse_resume_document(document)
    candidate_types = {candidate.candidate_type for candidate in candidates}
    skills = {
        str(candidate.payload["name"])
        for candidate in candidates
        if candidate.candidate_type == ResumeImportCandidateType.SKILL
    }
    headline = next(
        candidate
        for candidate in candidates
        if candidate.candidate_type == ResumeImportCandidateType.HEADLINE
    )
    experience = next(
        candidate
        for candidate in candidates
        if candidate.candidate_type == ResumeImportCandidateType.EXPERIENCE
    )

    assert document.quality in {"GOOD", "ACCEPTABLE"}
    assert document.quality_metrics["selected_modes"] == ["geometric"]
    assert any(block.text == "Pessoa Ficticia" for block in document.blocks)
    assert any(block.text == "Cidade Ficticia, Brasil" for block in document.blocks)
    assert ResumeImportCandidateType.SUMMARY in candidate_types
    assert ResumeImportCandidateType.EXPERIENCE in candidate_types
    assert ResumeImportCandidateType.PROJECT in candidate_types
    assert ResumeImportCandidateType.EDUCATION in candidate_types
    assert ResumeImportCandidateType.LANGUAGE in candidate_types
    assert {"Excel", "SQL", "Power BI", "Python", "Power Query"} <= skills
    assert headline.payload["headline"] == "Analista de Dados"
    assert "Pessoa Ficticia" not in str(headline.payload)
    assert "Cidade Ficticia" not in str(headline.payload)
    assert len(experience.block_ids) >= 4
    assert "Levantamento e analise" in str(experience.source_excerpt)
    assert (
        len(
            [
                candidate
                for candidate in candidates
                if candidate.candidate_type == ResumeImportCandidateType.AMBIGUOUS
            ]
        )
        <= 2
    )


def test_pdf_degraded_quality_gets_warning_and_low_confidence_only(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _patch_pdf_reader(
        monkeypatch,
        [
            _FakePdfPage(
                plain=_glued_resume_text(),
                layout=_glued_resume_text(),
                fragments=(),
            )
        ],
    )

    _upload, document = extract_resume("curriculo.pdf", _fake_pdf_bytes())
    candidates = parse_resume_document(document)

    assert document.quality == "DEGRADED"
    assert any("Tente o arquivo DOCX" in warning for warning in document.warnings)
    assert all(
        candidate.candidate_type != ResumeImportCandidateType.HEADLINE for candidate in candidates
    )
    assert all(
        candidate.confidence_label == ResumeImportConfidenceLabel.LOW for candidate in candidates
    )


def test_pdf_retry_updates_review_draft_with_same_file_hash(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    settings = _settings(tmp_path)
    run_migrations(settings)
    content = _fake_pdf_bytes()
    _patch_pdf_reader(
        monkeypatch,
        [
            _FakePdfPage(
                plain=_glued_resume_text(),
                layout=_glued_resume_text(),
                fragments=(),
            )
        ],
    )

    with session_scope(settings) as session:
        created = create_import_session(session, filename="curriculo.pdf", content=content)
        import_session = get_import_session(session, created.import_key)
        assert import_session.extraction_quality == "DEGRADED"
        assert {candidate.candidate_type for candidate in import_session.candidates} == {
            ResumeImportCandidateType.AMBIGUOUS
        }

        _patch_pdf_reader(
            monkeypatch,
            [
                _FakePdfPage(
                    plain=_glued_resume_text(),
                    layout=_glued_resume_text(),
                    fragments=_resume_fragments(),
                )
            ],
        )
        retried = retry_import_session(
            session,
            created.import_key,
            filename="curriculo.pdf",
            content=content,
            extraction_mode="geometric",
        )
        import_session = get_import_session(session, retried.import_key)

        assert import_session.extraction_quality in {"GOOD", "ACCEPTABLE"}
        assert import_session.extraction_mode == "geometric"
        assert any(
            candidate.candidate_type == ResumeImportCandidateType.EXPERIENCE
            for candidate in import_session.candidates
        )


def test_resume_session_requires_review_before_profile_and_confirms_with_provenance(
    tmp_path: Path,
) -> None:
    settings = _settings(tmp_path)
    run_migrations(settings)
    content = _resume_markdown()

    with session_scope(settings) as session:
        result = create_import_session(session, filename="curriculo.md", content=content)
        import_session = get_import_session(session, result.import_key)
        assert import_session.status == ResumeImportStatus.REVIEWING
        assert session.scalar(select(ProfessionalProfileVersion.id)) is None
        assert not any(
            "candidate@example.com" in c.original_payload_json for c in import_session.candidates
        )

        for candidate in import_session.candidates:
            if candidate.candidate_type in {
                ResumeImportCandidateType.SKILL,
                ResumeImportCandidateType.EXPERIENCE,
                ResumeImportCandidateType.PROJECT,
            }:
                accept_candidate(session, import_session.import_key, candidate.id)

        confirmed = confirm_import(session, import_session.import_key, activate_now=True)
        profile_version = session.get(
            ProfessionalProfileVersion, confirmed.profile.profile_version_id
        )
        assert profile_version is not None
        assert profile_version.is_active is True
        assert profile_version.source_format == "markdown"
        assert "resume_import" in profile_version.raw_profile_json
        assert "candidate@example.com" not in profile_version.raw_profile_json
        assert import_session.status == ResumeImportStatus.CONFIRMED
        assert import_session.confirmed_profile_version_id == profile_version.id

        repeated = confirm_import(session, import_session.import_key, activate_now=True)
        assert repeated.profile.profile_version_id == confirmed.profile.profile_version_id
        assert repeated.profile.created_version is False


def test_docx_import_can_be_discarded_without_profile(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    run_migrations(settings)

    with session_scope(settings) as session:
        created = create_import_session(session, filename="curriculo.docx", content=_resume_docx())
        import_session = discard_import(session, created.import_key)
        assert import_session.status == ResumeImportStatus.DISCARDED
        assert all(
            candidate.decision == ResumeImportDecision.PENDING
            for candidate in import_session.candidates
        )
        assert session.scalar(select(ProfessionalProfileVersion.id)) is None
        assert session.scalar(select(ResumeImportSession.id)) == import_session.id


def _resume_markdown() -> bytes:
    return b"""
# Perfil
Analista de dados junior

## Resumo
Estudante de dados com experiencia em dashboards e consultas SQL.

## Experiencia
Estagiario de Dados - Empresa X - 2024 - atual: Dashboards em Power BI, SQL e Python.

## Projetos
Dashboard de Vendas - Analise de indicadores com Python, SQL e Power BI.

## Habilidades
Python, SQL, Power BI

## Formacao
PUC Minas - Ciencia de Dados, cursando 2026

## Idiomas
Ingles intermediario

Contato: candidate@example.com
"""


def _resume_docx() -> bytes:
    document = Document()
    document.add_heading("Resumo", level=1)
    document.add_paragraph("Analista de dados com Python e SQL.")
    document.add_heading("Experiencia", level=1)
    document.add_paragraph(
        "Estagiario de Dados - Empresa X - 2024 - atual: Dashboards em Power BI."
    )
    document.add_heading("Habilidades", level=1)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Python"
    table.rows[0].cells[1].text = "SQL"
    content = BytesIO()
    document.save(content)
    return content.getvalue()


def _synthetic_pdf_text(items: Sequence[tuple[str, int, int, int]]) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    stream = DecodedStreamObject()
    stream.set_data(
        b"".join(
            (f"BT /F1 {font_size} Tf {x} {y} Td ({_pdf_escape(text)}) Tj ET\n").encode("latin-1")
            for text, x, y, font_size in items
        )
    )
    page[NameObject("/Contents")] = writer._add_object(stream)
    content = BytesIO()
    writer.write(content)
    return content.getvalue()


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _fake_pdf_bytes() -> bytes:
    return _synthetic_pdf_text([("PDF sintetico", 72, 720, 12)])


def _good_resume_text() -> str:
    return """
Analista de Dados
Resumo
Estudante de dados com experiencia em dashboards, SQL e Python.
Experiencia
Estagiaria de Dados - Empresa Sintetica - 2025 - atual
- Levantamento e analise de requisitos com mapeamento de fluxos de dados.
- Testes de APIs e validacao de dados com SQL e Python.
Projetos
Dashboard Academico - Painel em Power BI com Power Query e SQL.
Formacao
Universidade Ficticia - Analise e Desenvolvimento de Sistemas, cursando 2026
Habilidades
Excel, SQL, Power BI, Python e Power Query
Idiomas
Ingles intermediario
"""


def _glued_resume_text() -> str:
    return (
        "PessoaFicticiaCidadeFicticiaBrasil "
        "EmbuscadeoportunidadedeEstagioemAnalisededadosBIouAnalyticsaplicando"
        "ExcelSQLPowerBIPythonePowerQuerynaorganizacaonotratamentonaanalise"
        "ValidacaodedadosregrasdenegociotestesdeAPIseinvestigacaodeinconsistencias"
        "Formacaoembancosdedadosestruturasdedadosmodelagemdesistemasedesenvolvimento"
    )


def _resume_fragments() -> tuple[tuple[str, float, float, float], ...]:
    return (
        ("Pessoa", 72, 760, 12),
        ("Ficticia", 112, 760, 12),
        ("Cidade Ficticia, Brasil", 260, 760, 12),
        ("Analista de Dados", 72, 740, 13),
        ("Resumo", 72, 710, 14),
        ("Em busca de oportunidade de estagio em Analise de Dados, BI ou Analytics.", 72, 690, 12),
        ("Experiencia", 72, 660, 14),
        ("Estagiaria de Dados - Empresa Sintetica", 72, 640, 12),
        ("2025 - atual", 72, 622, 12),
        ("- Levantamento e analise de requisitos com mapeamento de fluxos de dados.", 86, 604, 12),
        ("- Testes de APIs e validacao de dados com SQL e Python.", 86, 586, 12),
        ("Projetos", 72, 556, 14),
        ("Dashboard Academico - Painel em Power BI com Power Query e SQL.", 72, 536, 12),
        ("Formacao", 72, 506, 14),
        ("Universidade Ficticia", 72, 486, 12),
        ("Analise e Desenvolvimento de Sistemas", 72, 468, 12),
        ("Cursando 2026", 72, 450, 12),
        ("Habilidades", 72, 420, 14),
        ("Excel, SQL, Power BI, Python e Power Query", 72, 400, 12),
        ("Idiomas", 72, 370, 14),
        ("Ingles intermediario", 72, 350, 12),
    )


class _FakePdfPage:
    def __init__(
        self,
        *,
        plain: str,
        layout: str,
        fragments: Sequence[tuple[str, float, float, float]],
    ) -> None:
        self.plain = plain
        self.layout = layout
        self.fragments = tuple(fragments)

    def extract_text(self, *args: Any, **kwargs: Any) -> str:
        _ = args
        visitor = kwargs.get("visitor_text")
        if callable(visitor):
            for text, x, y, font_size in self.fragments:
                visitor(text, (1, 0, 0, 1, 0, 0), (1, 0, 0, 1, x, y), {}, font_size)
            return ""
        if kwargs.get("extraction_mode") == "layout":
            return self.layout
        return self.plain


class _FakePdfReader:
    def __init__(self, pages: Sequence[_FakePdfPage]) -> None:
        self.is_encrypted = False
        self.pages = list(pages)


def _patch_pdf_reader(
    monkeypatch: pytest.MonkeyPatch,
    pages: Sequence[_FakePdfPage],
) -> None:
    import pypdf

    monkeypatch.setattr(pypdf, "PdfReader", lambda _content: _FakePdfReader(pages))


def _settings(tmp_path: Path) -> Settings:
    return Settings(
        database_url=f"sqlite:///{tmp_path / 'radar.sqlite3'}",
        config_dir=tmp_path / "config",
    )
